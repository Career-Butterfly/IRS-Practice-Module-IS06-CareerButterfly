import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import io
import os
import requests
from utils import resume_parser

def main():
    #set page to be wider
    st.set_page_config(layout='wide')
    st.title("Career Learning Hub")
    create_session_state()
    col1, col2 = st.columns([1, 5])
    with col1:
        back_button=st.button("Previous :arrow_left:")
        if back_button:
            st.switch_page("pages/job_recommender_page.py")
     
    with col2:
        #CV Generator       
        with st.expander("Cover Letter Generator"):
            with st.form(key='cv_form'):
                st.session_state.selected_job_company = st.text_input("Company Name: ", st.session_state.selected_job_company)
                st.session_state.selected_job_title = st.text_input("What role are you applying for? ", st.session_state.selected_job_title)
                st.session_state.contact_person = st.text_input("Who are you emailing? ", st.session_state.contact_person)
                st.session_state.resume_data["name"] = st.text_input("Name: ", st.session_state.resume_data["name"])
                st.session_state.my_email = st.text_input("Email: ", st.session_state.my_email)
                st.session_state.my_mobile_number = st.text_input("Mobile Number: ", st.session_state.my_mobile_number)
                st.session_state.resume_data["skill"] = st.text_input("Skills: ", st.session_state.resume_data["skill"])
                st.session_state.my_linkedin_profile = st.text_input("Linkedln Profile: ", st.session_state.my_linkedin_profile)
                st.session_state.personal_experience = st.text_input("I have experience in...", st.session_state.personal_experience)
                st.session_state.job_description = st.text_input("I am excited about the job because...", st.session_state.job_description)
                st.session_state.passion = st.text_input("I am passionate about...", st.session_state.passion)
                # job_specific = st.text_input("What do you like about this job? (Please keep this brief, one sentence only.) ")
                # specific_fit = st.text_input("Why do you think your experience is a good fit for this role? (Please keep this brief, one sentence only.) ")
                submit_button = st.form_submit_button(label='Generate Cover Letter')
            if submit_button:
                print("prompting cv generator service...")
                url = os.environ["CV_GENERATOR"]
                data = {"name": str(st.session_state.resume_data["name"]),
                        "email": str(st.session_state.my_email),
                        "mobile_number": str(st.session_state.my_mobile_number),
                        "linkedin_profile": str(st.session_state.my_linkedin_profile),
                        "skills": str(st.session_state.resume_data["skill"]),
                        "contact_person": str(st.session_state.contact_person),
                        "company_name": str(st.session_state.selected_job_company),
                        "role": str(st.session_state.selected_job_title),
                        "job_description": str(st.session_state.job_description),
                        "passion": str(st.session_state.passion)
                }
                response = requests.post(url, json=data)
                # Check if the request was successful (status code 200)
                if response.status_code == 200:
                    generated_cv = response.json()
                    st.session_state.generated_cv = generated_cv["generated_txt"]
            if st.session_state.generated_cv != '':
                st.subheader("Generated Cover Letter")
                st.write(st.session_state.generated_cv)
                bio = create_docx(str(st.session_state.generated_cv), "Cover Letter")
                st.download_button(label='Download Cover Letter', file_name=st.session_state.resume_data["name"]+'_cover_letter.docx', data=bio.getvalue(), mime="docx")
        #Resume Generator
        with st.expander("Resume Generator"):
            st.subheader("Generate a resume template based on the information you have filled in:")
            with st.form(key='resume_form'):
                submit_button = st.form_submit_button(label='Generate Resume')
            if submit_button:
                print("prompt resume generator service")
                url = os.environ["RESUME_GENERATOR"]
                data = {"name": str(st.session_state.resume_data["name"]),
                        "email": str(st.session_state.my_email),
                        "mobile_number": str(st.session_state.my_mobile_number),
                        "linkedin_profile": str(st.session_state.my_linkedin_profile),
                        "skills": str(st.session_state.resume_data["skill"]),
                        "education": str(resume_parser.retrieve_degree_info(st.session_state.resume_data)),
                        "works_experiences": str(resume_parser.retrieve_work_info(st.session_state.resume_data))
                }
                response = requests.post(url, json=data)
                # Check if the request was successful (status code 200)
                if response.status_code == 200:
                    generated_resume = response.json()
                    st.session_state.generated_resume = generated_resume["generated_txt"]
            if st.session_state.generated_resume != '':
                st.subheader("Generated Resume")
                st.write(st.session_state.generated_resume)
                bio = create_docx(str(st.session_state.generated_resume), "Resume")
                st.download_button(label='Download Resume', file_name=st.session_state.resume_data["name"]+'_resume.docx', data=bio.getvalue(), mime="docx")
        #Learning Resources Recommender
        with st.expander("Learning Resources"):
            with st.form(key='learning_resources'):
                submit_button = st.form_submit_button(label='generate learning resources')
            if submit_button:
                url = os.environ["LEARNING_RESOURCE_RECOMMENDER"]
                # Define the route endpoint and parameters
                route = '/generate_learning_resource_html_format'
                params = {
                    'param1': '',
                    'param2': st.session_state.selected_job_description,
                    'param3': st.session_state.selected_job_company
                }

                st.session_state.generated_learning_resources=requests.get(f"{url}{route}", params=params)
                
            if st.session_state.generated_learning_resources != '':
                st.components.v1.html(st.session_state.generated_learning_resources.text, scrolling=True, height=1080)
                st.download_button(
                    label="Download Learning Resource",
                    data=download_learning_resource_file(),
                    file_name="learning_resource.zip",
                    mime="application/zip"
                )
            
        st.text("Ready to start practicing for your interview?")
        go_to_interviewer_chatbot = st.button(label='I am ready!')
        ## switch to interviewer chatbot page
        if go_to_interviewer_chatbot:
            st.switch_page("pages/interview_chatbot_openai.py")
#initialize session state               
def create_session_state():
    session_state_keys=["role", "contact_person", "personal_experience", "job_description", "passion", "generated_cv", "generated_resume", "generated_learning_resources"]
    for key in session_state_keys:
        if key not in st.session_state:
            st.session_state[key] = ""
#create bio instance
def create_docx(generated_txt, header):
        document = Document()
        document.add_heading(header, level=0)
        document.add_paragraph(generated_txt)
        bio = io.BytesIO()
        document.save(bio)
        return bio
#download recommended learning resources as zip    
def download_learning_resource_file():
    # Define the base URL of your Flask app
    base_url = os.environ["LEARNING_RESOURCE_RECOMMENDER"]
    # Define the route endpoint and parameters
    route = '/download_learning_resource'
    # Make the GET request to the Flask route
    response = requests.get(f"{base_url}{route}")
    return response.content
                            
if __name__ == "__main__":
    main()